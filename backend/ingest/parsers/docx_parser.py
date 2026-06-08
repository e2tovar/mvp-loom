"""Parser DOCX: usa estilos de encabezado y señales de estilo/alineación para escenas.

research.md D4. Un separador de escena en .docx puede venir dado por el estilo o la
alineación del párrafo, no solo por símbolos visibles (FR-004a).
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.errors import InvalidFileError
from backend.ingest.parsers.base import Block, ParsedDocument, is_separator_line


def _heading_level(style_name: str) -> int | None:
    name = (style_name or "").strip().lower()
    if name.startswith("heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        return int(digits) if digits else 1
    if name in {"title", "título", "titulo"}:
        return 1
    return None


def _is_separator_style(style_name: str) -> bool:
    name = (style_name or "").strip().lower()
    return "separator" in name or "separador" in name


class DocxParser:
    def parse(self, path: Path) -> ParsedDocument:
        try:
            document = docx.Document(str(path))
        except Exception as exc:  # noqa: BLE001
            raise InvalidFileError(f"No se pudo leer el DOCX: {exc}") from exc

        blocks: list[Block] = []
        for para in document.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ""
            level = _heading_level(style_name)

            if level is not None and text:
                blocks.append(Block(kind="heading", text=text, level=level, style=style_name))
                continue

            # Señal de separador por estilo/alineación (propia de .docx) — FR-004a.
            centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
            centered_sep = centered and (not text or is_separator_line(text))
            if _is_separator_style(style_name) or centered_sep:
                blocks.append(Block(kind="separator", text=text or "***", style=style_name))
                continue

            if not text:
                continue
            if is_separator_line(text):
                blocks.append(Block(kind="separator", text=text, style=style_name))
            else:
                blocks.append(Block(kind="paragraph", text=text, style=style_name))

        title = (document.core_properties.title or "").strip() or None
        if not blocks:
            raise InvalidFileError("El DOCX no contiene texto extraíble.")
        return ParsedDocument(source_format="docx", blocks=blocks, title=title)
