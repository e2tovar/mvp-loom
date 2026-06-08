"""Genera de forma determinista las fixtures binarias (.epub, .docx) de M0.

Uso:  uv run python eval/fixtures/build_fixtures.py

Las fixtures de texto (.txt) y las anotaciones (*.annotation.json) se versionan
directamente; este builder solo produce los formatos binarios.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ebooklib import epub

HERE = Path(__file__).parent


def build_epub() -> Path:
    book = epub.EpubBook()
    book.set_identifier("crafted-two-chapters")
    book.set_title("Crafted Two Chapters")
    book.set_language("es")

    c1 = epub.EpubHtml(title="Capítulo 1", file_name="c1.xhtml", lang="es")
    c1.content = (
        "<html><body>"
        "<h1>Capítulo 1</h1>"
        "<p>Párrafo uno con acentós y eñes.</p>"
        "<p>Párrafo dos del primer capítulo.</p>"
        "<hr/>"
        "<p>Segunda escena tras el separador tipográfico.</p>"
        "</body></html>"
    )
    c2 = epub.EpubHtml(title="Capítulo 2", file_name="c2.xhtml", lang="es")
    c2.content = (
        "<html><body>"
        "<h1>Capítulo 2</h1>"
        "<p>El segundo capítulo no tiene separadores internos.</p>"
        "</body></html>"
    )
    book.add_item(c1)
    book.add_item(c2)
    book.toc = (c1, c2)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    # Spine sin 'nav' para que el orden de lectura sean solo los capítulos.
    book.spine = [c1, c2]

    out = HERE / "crafted-two-chapters.epub"
    epub.write_epub(str(out), book)
    return out


def build_docx() -> Path:
    document = docx.Document()
    document.core_properties.title = "Crafted Two Chapters"

    document.add_heading("Capítulo 1", level=1)
    document.add_paragraph("Párrafo uno con acentós y eñes.")
    document.add_paragraph("Párrafo dos del primer capítulo.")
    sep = document.add_paragraph("* * *")
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Segunda escena tras el separador.")

    document.add_heading("Capítulo 2", level=1)
    document.add_paragraph("El segundo capítulo no tiene separadores internos.")

    out = HERE / "crafted-two-chapters.docx"
    document.save(str(out))
    return out


if __name__ == "__main__":
    print("epub:", build_epub())
    print("docx:", build_docx())
